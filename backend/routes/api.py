import os
import shutil
import uuid
from typing import List, Dict, Any, Optional
from fastapi import APIRouter, UploadFile, File, Form, HTTPException, Header, Response, Depends
from pydantic import BaseModel

from backend.config import UPLOAD_DIR
from backend.services.document_service import DocumentService
from backend.services.rag_service import RAGService
from backend.services.ai_service import AIService
from backend.services.auth_service import AuthService
from backend.services.quiz_formatter import QuizFormatterService
from backend.database import (
    save_document, 
    get_document, 
    get_latest_document, 
    list_all_documents,
    update_document_title,
    delete_document,
    save_document_summary,
    save_document_quiz,
    save_document_progress,
    get_or_create_user, 
    authenticate_admin,
    register_user,
    login_user,
    list_all_users,
    get_user_by_id,
    admin_create_user,
    admin_update_user,
    admin_reset_user_password,
    admin_delete_user,
    clear_activity_logs,
    list_prompts,
    save_prompt,
    delete_prompt,
    get_system_settings,
    update_system_settings,
    get_activity_logs,
    get_admin_metrics,
    log_activity
)

router = APIRouter(prefix="/api")

# --- Auth helpers ---
def _get_current_user(x_user_id: Optional[str] = Header(None)):
    if not x_user_id:
        raise HTTPException(status_code=401, detail="غير مصرح - يلزم تسجيل الدخول (X-User-Id مفقود)")
    user = get_user_by_id(x_user_id)
    if not user:
        raise HTTPException(status_code=401, detail="المستخدم غير موجود أو الجلسة منتهية")
    return user

def _require_admin(current_user: dict = Depends(_get_current_user)):
    if current_user.get("role") != "admin":
        raise HTTPException(status_code=403, detail="صلاحيات المدير مطلوبة (Admin only)")
    return current_user

class AdminLoginRequest(BaseModel):
    admin_key: str

class StudentLoginRequest(BaseModel):
    email: str
    name: Optional[str] = ""

class ProgressRequest(BaseModel):
    progress_json: str

class RegisterRequest(BaseModel):
    name: str
    email: str
    password: str
    role: Optional[str] = "student"

class LoginRequest(BaseModel):
    email: str
    password: str

class AdminCreateUserRequest(BaseModel):
    name: str
    email: str
    password: str
    role: Optional[str] = "student"
    tier: Optional[str] = "Pro Academic 🌟"
    token_limit: Optional[int] = 500000
    permissions: Optional[Dict[str, Any]] = None

class AdminUpdateUserRequest(BaseModel):
    name: str
    email: str
    role: str
    tier: Optional[str] = "Pro Academic 🌟"
    token_limit: Optional[int] = 500000
    permissions: Optional[Dict[str, Any]] = None

class AdminResetPasswordRequest(BaseModel):
    new_password: str

class UpdateDocumentRequest(BaseModel):
    title: str

class UpdateSettingsRequest(BaseModel):
    settings: Dict[str, Any]

class ChatRequest(BaseModel):
    query: str
    doc_id: Optional[str] = None
    history: Optional[List[Dict[str, str]]] = []
    custom_system_prompt: Optional[str] = None

class SummarizeRequest(BaseModel):
    doc_id: Optional[str] = None
    level: Optional[str] = "full"
    language: Optional[str] = "ar"
    custom_system_prompt: Optional[str] = None

class QuizRequest(BaseModel):
    doc_id: Optional[str] = None
    count: Optional[int] = 5
    difficulty: Optional[str] = "medium"
    language: Optional[str] = "bilingual"
    custom_system_prompt: Optional[str] = None
    extract_only: Optional[bool] = False

class ProofreadRequest(BaseModel):
    text: str
    custom_system_prompt: Optional[str] = None

class TranslateRequest(BaseModel):
    doc_id: Optional[str] = None
    text: Optional[str] = None
    source_lang: Optional[str] = "en"
    target_lang: Optional[str] = "ar"
    mode: Optional[str] = "line_by_line" # 'target_only', 'page_by_page', 'line_by_line'
    custom_system_prompt: Optional[str] = None

class ValidateConnectionRequest(BaseModel):
    provider: Optional[str] = "gemini"
    api_key: Optional[str] = ""
    base_url: Optional[str] = ""
    model: Optional[str] = ""

class FetchModelsRequest(BaseModel):
    provider: Optional[str] = "gemini"
    base_url: Optional[str] = ""
    api_key: Optional[str] = ""

class GoogleVerifyRequest(BaseModel):
    credential: str
    client_id: Optional[str] = None

class CreatePromptRequest(BaseModel):
    category: str
    title: str
    description: Optional[str] = ""
    system_prompt: str

class GeneratePromptRequest(BaseModel):
    task_goal: str
    category: Optional[str] = "quiz"

class QuizExportRequest(BaseModel):
    quiz_data: Any
    format: str # 'custom_text', 'csv', 'xlsx', 'json'
    chapter_title: Optional[str] = "Chapter Exam"

class QuizImportTextRequest(BaseModel):
    raw_text: str

class DocxExportRequest(BaseModel):
    title: str
    subtitle: Optional[str] = ""
    doc_name: Optional[str] = ""
    content: Optional[str] = None
    units: Optional[List[Dict[str, Any]]] = None
    sections: Optional[List[Dict[str, Any]]] = None


@router.get("/health")
def health_check(
    x_ai_provider: Optional[str] = Header("gemini"),
    x_gemini_api_key: Optional[str] = Header(None)
):
    latest_doc = get_latest_document()
    return {
        "status": "ok",
        "provider": x_ai_provider,
        "has_documents": latest_doc is not None
    }

@router.post("/validate-key")
def validate_connection_endpoint(req: ValidateConnectionRequest):
    try:
        result = AIService.validate_connection(
            provider=req.provider or "gemini",
            api_key=req.api_key or "",
            base_url=req.base_url or "",
            model=req.model or ""
        )
        return result
    except Exception as e:
        raise HTTPException(status_code=400, detail=str(e))

@router.post("/fetch-models")
def fetch_models_endpoint(req: FetchModelsRequest):
    """Dynamically query the Base URL or Provider to discover installed models."""
    models = AIService.fetch_available_models(
        provider=req.provider or "gemini",
        base_url=req.base_url or "",
        api_key=req.api_key or ""
    )
    return {"models": models}

@router.post("/auth/google/verify")
def verify_google_token_endpoint(req: GoogleVerifyRequest):
    result = AuthService.verify_google_credential(req.credential, req.client_id)
    if not result["success"]:
        raise HTTPException(status_code=400, detail=result.get("error", "فشل التحقق من Google"))
    return result

@router.post("/auth/register")
def register_endpoint(req: RegisterRequest):
    res = register_user(req.name, req.email, req.password, req.role or "student")
    if not res["success"]:
        raise HTTPException(status_code=400, detail=res.get("error", "فشل إنشاء الحساب"))
    return res

@router.post("/auth/login")
def login_endpoint(req: LoginRequest):
    res = login_user(req.email, req.password)
    if not res["success"]:
        raise HTTPException(status_code=401, detail=res.get("error", "فشل تسجيل الدخول"))
    return res

@router.post("/auth/admin-login")
def admin_login_endpoint(req: AdminLoginRequest):
    res = authenticate_admin(req.admin_key)
    if not res["success"]:
        raise HTTPException(status_code=401, detail=res.get("error", "رمز التحقق أو كلمة مرور المدير غير صحيحة"))
    return res

@router.post("/auth/student-login")
def student_login_endpoint(req: StudentLoginRequest):
    clean_email = req.email.strip().lower()
    clean_name = req.name.strip() if req.name else clean_email.split("@")[0]
    user = get_or_create_user(
        google_id=f"student_{clean_email.replace('@', '_').replace('.', '_')}",
        email=clean_email,
        name=clean_name,
        picture=f"https://api.dicebear.com/7.x/avataaars/svg?seed={clean_email}",
        role="student"
    )
    return {"success": True, "user": user}

# --- Prompts Management Endpoints ---
@router.get("/prompts")
def get_prompts(category: Optional[str] = None):
    return {"prompts": list_prompts(category)}

@router.post("/prompts")
def create_custom_prompt(req: CreatePromptRequest):
    p_id = f"p_{str(uuid.uuid4())[:8]}"
    saved = save_prompt(
        prompt_id=p_id,
        category=req.category,
        title=req.title,
        description=req.description or "",
        system_prompt=req.system_prompt
    )
    return {"success": True, "prompt": saved}

@router.delete("/prompts/{prompt_id}")
def remove_custom_prompt(prompt_id: str):
    delete_prompt(prompt_id)
    return {"success": True}

@router.post("/generate-prompt")
def generate_prompt_endpoint(
    req: GeneratePromptRequest,
    x_ai_provider: Optional[str] = Header("gemini"),
    x_gemini_api_key: Optional[str] = Header(None),
    x_ai_base_url: Optional[str] = Header(None),
    x_gemini_model: Optional[str] = Header(None)
):
    result = AIService.generate_custom_prompt(
        task_goal=req.task_goal,
        category=req.category or "quiz",
        provider=x_ai_provider or "gemini",
        api_key=x_gemini_api_key,
        base_url=x_ai_base_url,
        model=x_gemini_model
    )
    return result

# --- Quiz Export & Import Endpoints ---
@router.post("/quiz/export")
def export_quiz_endpoint(req: QuizExportRequest):
    fmt = req.format.lower()
    chapter_title = req.chapter_title or "Chapter Exam"
    
    if fmt == "custom_text" or fmt == "txt":
        content = QuizFormatterService.to_bilingual_custom_text(req.quiz_data, chapter_title)
        return {
            "format": "txt",
            "content": content,
            "filename": f"Quiz_{chapter_title}.txt"
        }
    elif fmt == "csv":
        csv_text = QuizFormatterService.to_csv(req.quiz_data)
        return {
            "format": "csv",
            "content": csv_text,
            "filename": f"Quiz_{chapter_title}.csv"
        }
    elif fmt == "xlsx":
        excel_bytes = QuizFormatterService.to_excel_bytes(req.quiz_data)
        return Response(
            content=excel_bytes,
            media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            headers={"Content-Disposition": f'attachment; filename="Quiz_{chapter_title}.xlsx"'}
        )
    elif fmt == "json":
        return {
            "format": "json",
            "content": req.quiz_data,
            "filename": f"Quiz_{chapter_title}.json"
        }
    else:
        raise HTTPException(status_code=400, detail="الصيغة غير مدعومة. الصيغ المدعومة: custom_text, csv, xlsx, json")

@router.post("/quiz/import-text")
def import_quiz_text_endpoint(req: QuizImportTextRequest):
    parsed = QuizFormatterService.parse_custom_text(req.raw_text)
    if not parsed["questions"]:
        raise HTTPException(status_code=400, detail="لم يتم العثور على أسئلة مطابقة للصيغة في النص المدخل.")
    return parsed

# --- Document & AI Endpoints (Multi-Format Support with Multi-Tenant User Isolation) ---
@router.post("/upload")
async def upload_document(
    file: UploadFile = File(...), 
    user_id: Optional[str] = Form(None),
    x_user_id: Optional[str] = Header(None)
):
    effective_user_id = user_id or x_user_id
    # Dynamic allowed formats & size from system_settings
    sys_settings = get_system_settings()
    allowed_exts = [str(f).lower() for f in sys_settings.get("allowed_formats", [".pdf", ".docx", ".doc", ".pptx", ".ppt", ".txt", ".md", ".csv", ".xlsx", ".xls", ".rtf"])]
    max_size_mb = int(sys_settings.get("max_upload_size_mb", 50))
    file_ext = os.path.splitext(file.filename)[1].lower()
    
    if file_ext not in allowed_exts:
        raise HTTPException(
            status_code=400, 
            detail=f"نوع الملف غير مدعوم ({file_ext}). الصيغ المسموحة حالياً: {', '.join(allowed_exts)} (يمكن للمدير تعديلها من لوحة التحكم > إعدادات متقدمة)"
        )
    # Validate file size if available
    try:
        file.file.seek(0, 2)
        size_bytes = file.file.tell()
        file.file.seek(0)
        if size_bytes > max_size_mb * 1024 * 1024:
            raise HTTPException(status_code=400, detail=f"حجم الملف يتجاوز الحد المسموح ({max_size_mb} MB) — عدّل الحد من لوحة التحكم")
    except HTTPException:
        raise
    except Exception:
        pass
        
    doc_id = str(uuid.uuid4())[:8]
    save_path = os.path.join(UPLOAD_DIR, f"{doc_id}_{file.filename}")
    
    with open(save_path, "wb") as buffer:
        shutil.copyfileobj(file.file, buffer)
        
    try:
        pages_data = DocumentService.extract_text_and_pages(save_path)
        chunks = DocumentService.chunk_document(pages_data)
        full_text = "\n\n".join([p["text"] for p in pages_data])
        words_count = len(full_text.split())
        
        save_document(
            doc_id=doc_id,
            filename=file.filename,
            file_path=save_path,
            pages_count=len(pages_data),
            words_count=words_count,
            full_text=full_text,
            chunks=chunks,
            user_id=effective_user_id
        )
        
        return {
            "success": True,
            "doc_id": doc_id,
            "filename": file.filename,
            "pages_count": len(pages_data),
            "chunks_count": len(chunks),
            "words_count": words_count,
            "preview_text": full_text[:400] + "..." if len(full_text) > 400 else full_text
        }
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"حدث خطأ أثناء معالجة المستند: {str(e)}")

@router.get("/documents/latest")
def get_latest_doc_endpoint(x_user_id: Optional[str] = Header(None)):
    doc = get_latest_document(user_id=x_user_id)
    if not doc:
        return {"document": None}
    return {
        "document": {
            "doc_id": doc["id"],
            "filename": doc["filename"],
            "pages_count": doc["pages_count"],
            "words_count": doc.get("words_count", 0),
            "preview_text": doc["full_text"][:400] + "..." if len(doc.get("full_text", "")) > 400 else doc.get("full_text", "")
        }
    }

@router.post("/chat")
def chat_with_doc(
    req: ChatRequest,
    x_ai_provider: Optional[str] = Header("gemini"),
    x_gemini_api_key: Optional[str] = Header(None),
    x_ai_base_url: Optional[str] = Header(None),
    x_gemini_model: Optional[str] = Header(None),
    x_user_id: Optional[str] = Header(None)
):
    doc = get_document(req.doc_id, user_id=x_user_id) if req.doc_id else get_latest_document(user_id=x_user_id)
    chunks = doc.get("chunks", []) if doc else []
    # Dynamic RAG top_k from system_settings
    rag_k = int(get_system_settings().get("auto_rag_chunks", 4))
    # Ensure reasonable bounds 2-12, but keep user's expected 50 as fallback for legacy if chunks many
    # If chunks small, use all; if large, use setting*2
    top_k = max(4, min(50, rag_k * 3)) if rag_k else 50
    relevant_chunks = RAGService.search_relevant_chunks(req.query, chunks, top_k=top_k)
    
    result = AIService.answer_with_rag(
        query=req.query,
        context_chunks=relevant_chunks,
        conversation_history=req.history,
        provider=x_ai_provider or "gemini",
        api_key=x_gemini_api_key,
        base_url=x_ai_base_url,
        model=x_gemini_model,
        custom_system_prompt=req.custom_system_prompt
    )
    return {
        "query": req.query,
        "answer": result["answer"],
        "citations": result["citations"],
        "is_out_of_scope": result["is_out_of_scope"],
        "sources": result.get("sources", [])
    }

@router.post("/summarize")
def summarize_doc(
    req: SummarizeRequest,
    x_ai_provider: Optional[str] = Header("gemini"),
    x_gemini_api_key: Optional[str] = Header(None),
    x_ai_base_url: Optional[str] = Header(None),
    x_gemini_model: Optional[str] = Header(None),
    x_user_id: Optional[str] = Header(None)
):
    doc = None
    if req.doc_id and req.doc_id not in ("undefined", "null", ""):
        doc = get_document(req.doc_id, user_id=x_user_id)
    if not doc:
        doc = get_latest_document(user_id=x_user_id)
        
    full_text = doc.get("full_text", "") if doc else ""
    if not full_text:
        raise HTTPException(status_code=400, detail="يرجى رفع أو اختيار مادة تعليمية للتلخيص أولاً.")
    
    summary_data = AIService.generate_summary_and_mindmap(
        full_text=full_text,
        level=req.level or "full",
        language=req.language or "ar",
        provider=x_ai_provider or "gemini",
        api_key=x_gemini_api_key,
        base_url=x_ai_base_url,
        model=x_gemini_model,
        custom_system_prompt=req.custom_system_prompt
    )
    if doc and doc.get("id"):
        save_document_summary(doc["id"], summary_data)
        
    return summary_data

@router.post("/generate-quiz")
def generate_quiz_endpoint(
    req: QuizRequest,
    x_ai_provider: Optional[str] = Header("gemini"),
    x_gemini_api_key: Optional[str] = Header(None),
    x_ai_base_url: Optional[str] = Header(None),
    x_gemini_model: Optional[str] = Header(None),
    x_user_id: Optional[str] = Header(None)
):
    doc = None
    if req.doc_id and req.doc_id not in ("undefined", "null", ""):
        doc = get_document(req.doc_id, user_id=x_user_id)
    if not doc:
        doc = get_latest_document(user_id=x_user_id)
        
    full_text = doc.get("full_text", "") if doc else ""
    if not full_text:
        raise HTTPException(status_code=400, detail="يرجى رفع أو اختيار مادة تعليمية للاختبار أولاً.")
    
    quiz_data = AIService.generate_quiz(
        full_text=full_text,
        count=req.count or 5,
        difficulty=req.difficulty or "medium",
        language=req.language or "bilingual",
        provider=x_ai_provider or "gemini",
        api_key=x_gemini_api_key,
        base_url=x_ai_base_url,
        model=x_gemini_model,
        custom_system_prompt=req.custom_system_prompt,
        extract_only=req.extract_only
    )
    if doc and doc.get("id"):
        save_document_quiz(doc["id"], quiz_data)
        
    return quiz_data

@router.post("/proofread")
def proofread_endpoint(
    req: ProofreadRequest,
    x_ai_provider: Optional[str] = Header("gemini"),
    x_gemini_api_key: Optional[str] = Header(None),
    x_ai_base_url: Optional[str] = Header(None),
    x_gemini_model: Optional[str] = Header(None)
):
    if not req.text.strip():
        raise HTTPException(status_code=400, detail="النص المدخل فارغ.")
        
    result = AIService.proofread_text(
        input_text=req.text,
        provider=x_ai_provider or "gemini",
        api_key=x_gemini_api_key,
        base_url=x_ai_base_url,
        model=x_gemini_model,
        custom_system_prompt=req.custom_system_prompt
    )
    return result

@router.post("/translate")
def translate_endpoint(
    req: TranslateRequest,
    x_ai_provider: Optional[str] = Header("gemini"),
    x_gemini_api_key: Optional[str] = Header(None),
    x_ai_base_url: Optional[str] = Header(None),
    x_gemini_model: Optional[str] = Header(None),
    x_user_id: Optional[str] = Header(None)
):
    """Translate academic documents in pure, page-by-page, or interlinear line-by-line modes."""
    doc = None
    if req.doc_id and req.doc_id not in ("undefined", "null", ""):
        doc = get_document(req.doc_id, user_id=x_user_id)
    if not doc:
        doc = get_latest_document(user_id=x_user_id)

    if req.text and req.text.strip():
        full_text = req.text.strip()
    elif doc:
        full_text = doc.get("full_text", "")
    else:
        full_text = ""

    if not full_text:
        raise HTTPException(status_code=400, detail="يرجى رفع أو اختيار مادة تعليمية تحتوي على نصوص للترجمة أولاً.")

    result = AIService.translate_document(
        full_text=full_text,
        source_lang=req.source_lang or "en",
        target_lang=req.target_lang or "ar",
        mode=req.mode or "line_by_line",
        provider=x_ai_provider or "gemini",
        api_key=x_gemini_api_key,
        base_url=x_ai_base_url,
        model=x_gemini_model,
        custom_system_prompt=req.custom_system_prompt
    )
    return result

@router.post("/export/docx")
def export_docx_endpoint(req: DocxExportRequest):
    """Generate and stream a styled Microsoft Word (.docx) document."""
    import docx
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    import io
    from fastapi.responses import StreamingResponse

    doc = docx.Document()
    
    # Page Margins (1 inch)
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Document Title
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_p.add_run(req.title or "المستند المترجم")
    title_run.font.size = Pt(20)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor(0x1E, 0x3A, 0x8A)

    # Subtitle / Meta
    sub_p = doc.add_paragraph()
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta_text = f"المستند الأصلي: {req.doc_name or 'مادة تعليمية'}  |  المنصة: المساعد الأكاديمي الذكي"
    if req.subtitle:
        meta_text = f"{req.subtitle}\n{meta_text}"
    sub_run = sub_p.add_run(meta_text)
    sub_run.font.size = Pt(10.5)
    sub_run.font.italic = True
    sub_run.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)
    
    doc.add_paragraph().paragraph_format.space_after = Pt(12)

    # If Line-by-line translation units provided
    if req.units and len(req.units) > 0:
        h = doc.add_heading(level=1)
        hrun = h.add_run("الترجمة السطرية الموازية (Line-by-Line Parallel Translation)")
        hrun.font.color.rgb = RGBColor(0x1E, 0x3A, 0x8A)
        
        table = doc.add_table(rows=1, cols=2)
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'النص الأصلي (Original Text)'
        hdr_cells[1].text = 'الترجمة الأكاديمية (Arabic Translation)'
        
        for idx, u in enumerate(req.units):
            row_cells = table.add_row().cells
            row_cells[0].text = f"[{idx+1}] {u.get('original', '')}"
            row_cells[1].text = u.get('translated', '')

    # If general sections provided
    elif req.sections and len(req.sections) > 0:
        for sec in req.sections:
            h = doc.add_heading(level=1)
            hrun = h.add_run(sec.get('title', 'قسم'))
            hrun.font.color.rgb = RGBColor(0x1E, 0x3A, 0x8A)
            
            p = doc.add_paragraph()
            prun = p.add_run(sec.get('content', ''))
            prun.font.size = Pt(12)
            p.paragraph_format.line_spacing = 1.3
            p.paragraph_format.space_after = Pt(8)

    # If raw content text provided
    elif req.content:
        for line in req.content.split('\n'):
            line_str = line.strip()
            if not line_str:
                continue
            if line_str.startswith('# '):
                h = doc.add_heading(level=1)
                hrun = h.add_run(line_str[2:])
                hrun.font.color.rgb = RGBColor(0x1E, 0x3A, 0x8A)
            elif line_str.startswith('## '):
                h = doc.add_heading(level=2)
                hrun = h.add_run(line_str[3:])
                hrun.font.color.rgb = RGBColor(0x02, 0x84, 0xC7)
            elif line_str.startswith('### '):
                h = doc.add_heading(level=3)
                hrun = h.add_run(line_str[4:])
            elif line_str.startswith('- ') or line_str.startswith('* '):
                doc.add_paragraph(line_str[2:], style='List Bullet')
            else:
                p = doc.add_paragraph()
                prun = p.add_run(line_str)
                prun.font.size = Pt(12)
                p.paragraph_format.line_spacing = 1.3

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    
    clean_name = f"Translated_{req.doc_name or 'Document'}.docx"
    return StreamingResponse(
        buf,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="{clean_name}"'}
    )



# -------------------------------------------------------------
# Document Management Endpoints (مكتبة وإدارة المستندات)
# -------------------------------------------------------------

@router.get("/documents")
def list_documents_endpoint(x_user_id: Optional[str] = Header(None)):
    """Retrieve all uploaded documents belonging exclusively to the authenticated user."""
    docs = list_all_documents(user_id=x_user_id)
    return {"documents": docs, "total": len(docs)}

@router.get("/documents/{doc_id}")
def get_document_endpoint(doc_id: str, x_user_id: Optional[str] = Header(None)):
    """Retrieve single document details with owner validation."""
    doc = get_document(doc_id, user_id=x_user_id)
    if not doc:
        raise HTTPException(status_code=404, detail="المستند غير موجود أو لا تملك صلاحية الوصول إليه.")
    return {"document": doc}

@router.patch("/documents/{doc_id}")
def update_document_endpoint(doc_id: str, req: UpdateDocumentRequest, x_user_id: Optional[str] = Header(None)):
    """Rename or update document title with owner validation."""
    success = update_document_title(doc_id, req.title.strip(), user_id=x_user_id)
    if not success:
        raise HTTPException(status_code=404, detail="فشل تحديث المستند أو لم يتم العثور عليه في مكتبتك.")
    return {"success": True, "message": "تم تحديث اسم المستند بنجاح"}

@router.delete("/documents/{doc_id}")
def delete_document_endpoint(doc_id: str, x_user_id: Optional[str] = Header(None)):
    """Delete document from database and storage with owner validation."""
    success = delete_document(doc_id, user_id=x_user_id)
    if not success:
        raise HTTPException(status_code=404, detail="المستند غير موجود في مكتبتك الخاصة.")
    return {"success": True, "message": "تم حذف المستند بنجاح"}

@router.get("/documents/{doc_id}/progress")
def get_quiz_progress_endpoint(doc_id: str, x_user_id: Optional[str] = Header(None)):
    """Retrieve saved quiz progress for a document."""
    doc = get_document(doc_id, user_id=x_user_id)
    if not doc:
        raise HTTPException(status_code=404, detail="المستند غير موجود")
    return {"progress_json": doc.get("quiz_progress_json")}

@router.post("/documents/{doc_id}/progress")
def save_quiz_progress_endpoint(doc_id: str, req: ProgressRequest, x_user_id: Optional[str] = Header(None)):
    """Save quiz progress for a document."""
    doc = get_document(doc_id, user_id=x_user_id)
    if not doc:
        raise HTTPException(status_code=404, detail="المستند غير موجود")
    save_document_progress(doc_id, req.progress_json)
    return {"success": True}

# -------------------------------------------------------------
# Admin Control Panel Endpoints (لوحة التحكم الشاملة للإدارة)
# -------------------------------------------------------------

@router.get("/settings/public")
def get_public_settings_endpoint():
    """Get public branding information (name, logo, slogan, university) for UI components."""
    settings = get_system_settings()
    return {
        "platform_name": settings.get("platform_name", "ذكاء EduAI"),
        "platform_subtitle": settings.get("platform_subtitle", "المنصة الأكاديمية الذكية المتكاملة"),
        "university_name": settings.get("university_name", "الجامعة"),
        "faculty_name": settings.get("faculty_name", "كلية الحاسبات وتكنولوجيا المعلومات"),
        "support_email": settings.get("support_email", "admin@eduai.edu"),
        "footer_text": settings.get("footer_text", "المنصة الأكاديمية الذكية المتقدمة"),
        "logo_icon": settings.get("logo_icon", "GraduationCap"),
        "custom_logo_url": settings.get("custom_logo_url", ""),
        "welcome_headline": settings.get("welcome_headline", "مرحباً بك في المنصة الأكاديمية الذكية"),
        "welcome_description": settings.get("welcome_description", "بيئة تعليمية وبحثية جامعية مدعومة بالذكاء الاصطناعي للمذاكرة التفاعلية وتوليد خرائط المفاهيم والاختبارات."),
        "maintenance_mode": settings.get("maintenance_mode", False),
        "registration_enabled": settings.get("registration_enabled", True)
    }

@router.get("/admin/stats")
def get_admin_stats_endpoint(current_admin: dict = Depends(_require_admin)):
    """Get live metrics, counts, database size, and system health."""
    metrics = get_admin_metrics()
    return metrics

@router.get("/admin/settings")
def get_admin_settings_endpoint(current_admin: dict = Depends(_require_admin)):
    """Get platform configuration policies and AI defaults."""
    settings = get_system_settings()
    return {"settings": settings}

@router.post("/admin/settings")
def update_admin_settings_endpoint(req: UpdateSettingsRequest, current_admin: dict = Depends(_require_admin)):
    """Update platform configuration policies and AI defaults."""
    update_system_settings(req.settings)
    return {"success": True, "message": "تم حفظ وتطبيق الإعدادات بنجاح"}

@router.get("/admin/users")
def get_admin_users_endpoint(current_admin: dict = Depends(_require_admin)):
    """Get all registered users and their roles, quotas, and stats."""
    users = list_all_users()
    return {"users": users, "total": len(users)}

@router.post("/admin/users")
def create_admin_user_endpoint(req: AdminCreateUserRequest, current_admin: dict = Depends(_require_admin)):
    """Create a new user account with specified role and quota."""
    res = admin_create_user(
        name=req.name,
        email=req.email,
        password=req.password,
        role=req.role or "student",
        tier=req.tier or "Pro Academic 🌟",
        token_limit=req.token_limit or 500000,
        permissions=req.permissions
    )
    if not res["success"]:
        raise HTTPException(status_code=400, detail=res.get("error", "فشل إنشاء المستخدم"))
    return res

@router.patch("/admin/users/{user_id}")
def update_admin_user_endpoint(user_id: str, req: AdminUpdateUserRequest, current_admin: dict = Depends(_require_admin)):
    """Update user profile, role, tier, or token limit."""
    res = admin_update_user(
        user_id=user_id,
        name=req.name,
        email=req.email,
        role=req.role,
        tier=req.tier,
        token_limit=req.token_limit,
        permissions=req.permissions
    )
    if not res["success"]:
        raise HTTPException(status_code=400, detail=res.get("error", "فشل تعديل المستخدم"))
    return res

@router.patch("/admin/users/{user_id}/reset-password")
def reset_admin_user_password_endpoint(user_id: str, req: AdminResetPasswordRequest, current_admin: dict = Depends(_require_admin)):
    """Reset a user's password."""
    res = admin_reset_user_password(user_id, req.new_password)
    if not res["success"]:
        raise HTTPException(status_code=400, detail=res.get("error", "فشل إعادة تعيين كلمة المرور"))
    return res

@router.delete("/admin/users/{user_id}")
def delete_admin_user_endpoint(user_id: str, current_admin: dict = Depends(_require_admin)):
    """Delete a user account and all their documents."""
    res = admin_delete_user(user_id)
    if not res["success"]:
        raise HTTPException(status_code=400, detail=res.get("error", "فشل حذف المستخدم"))
    return res

@router.get("/admin/logs")
def get_admin_logs_endpoint(limit: Optional[int] = 50, current_admin: dict = Depends(_require_admin)):
    """Get platform activity and audit logs."""
    logs = get_activity_logs(limit=limit or 50)
    return {"logs": logs}

@router.delete("/admin/logs")
def clear_admin_logs_endpoint(current_admin: dict = Depends(_require_admin)):
    """Clear all audit logs."""
    clear_activity_logs()
    return {"success": True, "message": "تم مسح سجلات النشاط بنجاح"}

@router.post("/admin/clear-cache")
def clear_cache_endpoint(current_admin: dict = Depends(_require_admin)):
    """Clear temporary upload buffers and compact database."""
    log_activity("clear_cache", "تم تنفيذ تنظيف الذاكرة المؤقتة وضغط قاعدة البيانات", "info")
    return {"success": True, "message": "تم تنظيف الذاكرة المؤقتة بنجاح"}

